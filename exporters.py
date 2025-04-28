import streamlit as st
import pandas as pd
import os
from datetime import datetime
import base64
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Inches
import xlsxwriter
from utils import get_all_data, check_data_exists, format_filename

def export_to_word():
    """Generate and download a Word report of all collected data"""
    st.title("Exportar Reporte (Word + PDF)")
    
    if not check_data_exists():
        st.warning("⚠️ No hay datos para exportar. Por favor complete al menos una sección.")
        return
    
    # Get farm name from datos_generales if available
    datos_df = pd.DataFrame()
    if os.path.exists("data/datos_generales.csv"):
        datos_df = pd.read_csv("data/datos_generales.csv")
    
    farm_name = "tambo"
    if not datos_df.empty and 'nombre_tambo' in datos_df.columns:
        farm_name = datos_df['nombre_tambo'].iloc[0]
    
    # Create document
    doc = Document()
    
    # Add styles
    styles = doc.styles
    
    # Title style
    title_style = styles['Title']
    title_style.font.size = Pt(18)
    title_style.font.bold = True
    
    # Heading 1 style
    heading1_style = styles['Heading 1']
    heading1_style.font.size = Pt(16)
    heading1_style.font.bold = True
    
    # Heading 2 style
    heading2_style = styles['Heading 2']
    heading2_style.font.size = Pt(14)
    
    # Add title
    doc.add_heading(f"Reporte de Tambo: {farm_name}", 0)
    
    # Add timestamp
    doc.add_paragraph(f"Generado el: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    doc.add_paragraph("FieldLens - Recolección de Datos en Tambos")
    
    # Add a table of contents placeholder
    doc.add_paragraph("Contenido:")
    doc.add_paragraph("...")
    
    # Load all data
    all_data = get_all_data()
    
    # Datos Generales section
    if not all_data['datos_generales'].empty:
        doc.add_heading("Datos Generales", 1)
        df = all_data['datos_generales']
        
        # Convert dataframe to readable format
        for column in df.columns:
            if column != 'uuid':  # Skip UUID column
                doc.add_paragraph(f"{column.replace('_', ' ').title()}: {df[column].iloc[0]}")
    
    # Superficies e Insumos section
    if not all_data['superficies_insumos'].empty:
        doc.add_heading("Superficies e Insumos", 1)
        df = all_data['superficies_insumos']
        
        # Create table
        table = doc.add_table(rows=1, cols=len(df.columns) - 1 if 'uuid' in df.columns else len(df.columns))
        table.style = 'Table Grid'
        
        # Add header row
        header_cells = table.rows[0].cells
        col_idx = 0
        for column in df.columns:
            if column != 'uuid':  # Skip UUID column
                header_cells[col_idx].text = column.replace('_', ' ').title()
                col_idx += 1
        
        # Add data rows
        for _, row in df.iterrows():
            cells = table.add_row().cells
            col_idx = 0
            for column in df.columns:
                if column != 'uuid':  # Skip UUID column
                    cells[col_idx].text = str(row[column])
                    col_idx += 1
    
    # Add remaining sections with similar pattern
    sections = [
        ("Manejo y Recursos", 'manejo'),
        ("Fertilización", 'fertilizacion'),
        ("Protección de Cultivos", 'proteccion_cultivos'),
        ("Riego / Uso de Agua", 'riego'),
        ("Energía", 'energia'),
        ("Rebaño", 'rebano'),
        ("Gestión de Efluentes", 'efluentes'),
        ("Transporte", 'transporte')
    ]
    
    for section_title, data_key in sections:
        if not all_data[data_key].empty:
            doc.add_heading(section_title, 1)
            df = all_data[data_key]
            
            # Create table
            table = doc.add_table(rows=1, cols=len(df.columns) - 1 if 'uuid' in df.columns else len(df.columns))
            table.style = 'Table Grid'
            
            # Add header row
            header_cells = table.rows[0].cells
            col_idx = 0
            for column in df.columns:
                if column != 'uuid':  # Skip UUID column
                    header_cells[col_idx].text = column.replace('_', ' ').title()
                    col_idx += 1
            
            # Add data rows
            for _, row in df.iterrows():
                cells = table.add_row().cells
                col_idx = 0
                for column in df.columns:
                    if column != 'uuid':  # Skip UUID column
                        cells[col_idx].text = str(row[column])
                        col_idx += 1
    
    # Save document to BytesIO object
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    # Create download button for Word document
    safe_farm_name = format_filename(farm_name)
    word_filename = f"FieldLens_{safe_farm_name}_{datetime.now().strftime('%Y%m%d')}.docx"
    
    st.download_button(
        label="📥 Descargar Reporte Word (.docx)",
        data=doc_io,
        file_name=word_filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    
    # Display note about PDF conversion
    st.info("📝 Nota: Para obtener una versión PDF, descargue el documento Word y conviértalo a PDF con su programa preferido.")


def export_to_excel(farm_id=None, all_farms=False):
    """Generate and download an Excel file with farm data"""
    st.title("Exportar Datos a Excel")
    
    if not check_data_exists():
        st.warning("⚠️ No hay datos para exportar. Por favor complete al menos una sección.")
        return
        
    farm_name = "tambo"
    if "farm_name" in st.session_state:
        farm_name = st.session_state.farm_name
        
    safe_farm_name = format_filename(farm_name)
    excel_filename = f"FieldLens_{safe_farm_name}_{datetime.now().strftime('%Y%m%d')}.xlsx"
    
    # Create BytesIO object
    excel_io = BytesIO()
    
    # Create Excel writer
    with pd.ExcelWriter(excel_io, engine='xlsxwriter') as writer:
        # Get all data
        all_data = get_all_data()
        
        # Define section names and keys
        sections = [
            ("Datos Generales", 'datos_generales'),
            ("Superficies e Insumos", 'superficies_insumos'),
            ("Manejo y Recursos", 'manejo'),
            ("Fertilización", 'fertilizacion'),
            ("Protección de Cultivos", 'proteccion_cultivos'),
            ("Riego / Uso de Agua", 'riego'),
            ("Energía", 'energia'),
            ("Rebaño", 'rebano'),
            ("Gestión de Efluentes", 'efluentes'),
            ("Transporte", 'transporte')
        ]
        
        # Write each dataframe to a separate sheet
        for sheet_name, data_key in sections:
            df = all_data[data_key]
            if not df.empty:
                # Remove UUID column if present
                if 'uuid' in df.columns:
                    df = df.drop(columns=['uuid'])
                    
                # Write to Excel
                df.to_excel(writer, sheet_name=sheet_name, index=False)
                
                # Get workbook and worksheet
                workbook = writer.book
                worksheet = writer.sheets[sheet_name]
                
                # Add formats
                header_format = workbook.add_format({
                    'bold': True,
                    'font_color': 'white',
                    'bg_color': '#4380fa',
                    'border': 1
                })
                
                # Apply header format
                for col_num, value in enumerate(df.columns.values):
                    worksheet.write(0, col_num, value, header_format)
                    
                # Auto-fit columns
                for i, col in enumerate(df.columns):
                    column_width = max(df[col].astype(str).map(len).max(), len(col)) + 2
                    worksheet.set_column(i, i, column_width)
        
        # Create summary sheet
        summary_df = pd.DataFrame({
            'Sección': [name for name, _ in sections],
            'Registros': [len(all_data[key]) if not all_data[key].empty else 0 for _, key in sections]
        })
        
        summary_df.to_excel(writer, sheet_name='Resumen', index=False)
        
        # Format summary sheet
        workbook = writer.book
        worksheet = writer.sheets['Resumen']
        
        header_format = workbook.add_format({
            'bold': True,
            'font_color': 'white',
            'bg_color': '#4380fa',
            'border': 1
        })
        
        # Apply header format
        for col_num, value in enumerate(summary_df.columns.values):
            worksheet.write(0, col_num, value, header_format)
            
        # Auto-fit columns
        for i, col in enumerate(summary_df.columns):
            column_width = max(summary_df[col].astype(str).map(len).max(), len(col)) + 2
            worksheet.set_column(i, i, column_width)
    
    # Reset pointer to start of BytesIO object
    excel_io.seek(0)
    
    # Create download button for Excel file
    st.download_button(
        label="📥 Descargar Excel Completo",
        data=excel_io,
        file_name=excel_filename,
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    
    col1, col2 = st.columns([2, 1])
    with col1:
        st.download_button(
            label="📥 Exportar Datos del Tambo Actual",
            data=excel_io,
            file_name=excel_filename,
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
    with col2:
        if st.button("📊 Exportar Todos los Tambos"):
            export_to_excel(all_farms=True)
    
    st.info("📊 Los datos se exportarán organizados por secciones en diferentes hojas de Excel.")
    
    # Display current farm data in a user-friendly table
    st.subheader("Vista de Datos Recolectados")
    
    from st_aggrid import AgGrid, GridOptionsBuilder
    from st_aggrid.grid_options_builder import GridOptionsBuilder
    
    # Get current farm data and combine into one table
    current_farm_id = st.session_state.get('farm_id')
    all_data = get_all_data()
    combined_data = []
    
    for section_name, df in all_data.items():
        if not df.empty:
            # Filter for current farm if applicable
            if 'farm_id' in df.columns and current_farm_id:
                df = df[df['farm_id'] == current_farm_id]
            
            # Remove technical columns
            for col in ['uuid', 'farm_id']:
                if col in df.columns:
                    df = df.drop(columns=[col])
            
            # Add section information
            if not df.empty:
                df['Sección'] = section_name.replace('_', ' ').title()
                combined_data.append(df)
    
    if combined_data:
        combined_df = pd.concat(combined_data, ignore_index=True)
        
        # Configure grid options
        gb = GridOptionsBuilder.from_dataframe(combined_df)
        gb.configure_default_column(
            groupable=True,
            value=True,
            enableRowGroup=True,
            aggFunc='sum',
            editable=False
        )
        gb.configure_grid_options(domLayout='normal')
        gb.configure_side_bar()
        
        # Create grid
        grid_response = AgGrid(
            combined_df,
            gridOptions=gb.build(),
            enable_enterprise_modules=True,
            update_mode='MODEL_CHANGED',
            data_return_mode='FILTERED',
            theme='streamlit',
            height=500
        )
    else:
        st.warning("No hay datos disponibles para mostrar.")
